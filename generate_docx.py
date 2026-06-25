"""
generate_docx.py
Generates Assignment_1_Solution.docx using python-docx.
Run: pip install python-docx && python generate_docx.py
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import os
except ImportError:
    print("Installing python-docx...")
    import subprocess, sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Styles helper ─────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, bold=False, italic=False, size=11,
             align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_formula(doc, formula_text):
    """Add a formula line in monospace, centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(formula_text)
    run.font.name = "Courier New"
    run.font.size = Pt(11)
    run.bold = True
    return p

def add_separator(doc):
    p = doc.add_paragraph("─" * 75)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.runs[0]
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(180, 180, 180)

def add_answer(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("✅ Final Answer: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0, 100, 0)
    return p

# ─────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RTC INSTITUTE OF TECHNOLOGY")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(26, 26, 78)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Anandi, Ormanjhi, Ranchi – 835219, Jharkhand")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(80, 80, 80)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Department of Computer Science & Engineering")
run.bold = True
run.font.size = Pt(13)

add_separator(doc)
doc.add_paragraph()

# Cover info table
table = doc.add_table(rows=10, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
fields = [
    ("Assignment No.",       "01"),
    ("Course Name",          "Data Analysis using Open-Source Tools"),
    ("Submitted To",         "Mr. Basudev Mahata, Dept. CSE, RTCIT Ranchi"),
    ("Name of Student",      "Bablu Kumar"),
    ("University Roll No.",  "23052440012"),
    ("College Roll No.",     "23011036"),
    ("Session",              "2025-26"),
    ("Semester",             "6th"),
    ("Batch",                "2023-27"),
    ("Date of Submission",   "09/05/26"),
]
for i,(k,v) in enumerate(fields):
    r = table.rows[i]
    r.cells[0].text = k
    r.cells[1].text = v
    r.cells[0].paragraphs[0].runs[0].bold = True
    r.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    r.cells[1].paragraphs[0].runs[0].font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Assignment 1 — Statistics, Probability & Python Foundations")
run.bold = True
run.font.size = Pt(15)
run.font.color.rgb = RGBColor(26, 26, 78)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# TASK 1
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "TASK 1 — Descriptive Statistics, Probability & Combinatorics", level=1, color=(26,26,78))

# ── Q 1.1 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "Question 1.1 — Descriptive Statistics", level=2)
add_para(doc, "Given Marks: 72, 85, 90, 65, 78, 92, 68, 75, 88, 80", bold=True)
add_para(doc, "n = 10 observations", bold=True)
doc.add_paragraph()

# MEAN
add_para(doc, "1. MEAN", bold=True, size=12, color=(26,26,78))
add_para(doc, "Formula:")
add_formula(doc, "x_bar = (Sum of all values) / n")
add_para(doc, "Substitution:")
add_formula(doc, "x_bar = (72 + 85 + 90 + 65 + 78 + 92 + 68 + 75 + 88 + 80) / 10")
add_para(doc, "Intermediate Calculation:")
add_formula(doc, "x_bar = 793 / 10")
add_answer(doc, "Mean = 79.3")
add_para(doc, "Explanation: The mean is the arithmetic average. Sum = 793, divided by 10 gives 79.3. The average student scored 79.3 marks.", italic=True, color=(80,80,80))

add_separator(doc)

# MEDIAN
add_para(doc, "2. MEDIAN", bold=True, size=12, color=(26,26,78))
add_para(doc, "Formula: For n even: Median = (x(n/2) + x(n/2+1)) / 2")
add_para(doc, "Step 1 — Sort in ascending order:")
add_formula(doc, "65, 68, 72, 75, 78, 80, 85, 88, 90, 92")
add_para(doc, "Step 2 — n = 10 (even), 5th value = 78, 6th value = 80:")
add_formula(doc, "Median = (78 + 80) / 2 = 158 / 2")
add_answer(doc, "Median = 79.0")
add_para(doc, "Explanation: The median divides the dataset in half. 50% scored below 79 and 50% scored above 79.", italic=True, color=(80,80,80))

add_separator(doc)

# MODE
add_para(doc, "3. MODE", bold=True, size=12, color=(26,26,78))
add_para(doc, "Definition: Value appearing most frequently.")
add_para(doc, "Frequency count of each value: All values appear exactly once (frequency = 1).")
add_answer(doc, "No Mode exists — all values are unique.")
add_para(doc, "Explanation: Since every mark appears only once, there is no mode. This is a uniform frequency distribution.", italic=True, color=(80,80,80))

add_separator(doc)

# RANGE
add_para(doc, "4. RANGE", bold=True, size=12, color=(26,26,78))
add_para(doc, "Formula:")
add_formula(doc, "Range = Maximum - Minimum")
add_para(doc, "Substitution:")
add_formula(doc, "Range = 92 - 65")
add_answer(doc, "Range = 27")
add_para(doc, "Explanation: The range measures total spread. The 27-mark difference shows moderate variability in the data.", italic=True, color=(80,80,80))

add_separator(doc)

# VARIANCE
add_para(doc, "5. SAMPLE VARIANCE (n-1)", bold=True, size=12, color=(26,26,78))
add_para(doc, "Formula:")
add_formula(doc, "s^2 = Sum[(xi - x_bar)^2] / (n-1)")
add_para(doc, "Given: x_bar = 79.3, n = 10")
add_para(doc, "Step 1 — Calculate each (xi - x_bar)^2:")

rows_data = [
    ("xi","xi - x_bar","(xi - x_bar)^2"),
    ("72","72-79.3 = -7.3","53.29"),
    ("85","85-79.3 = +5.7","32.49"),
    ("90","90-79.3 = +10.7","114.49"),
    ("65","65-79.3 = -14.3","204.49"),
    ("78","78-79.3 = -1.3","1.69"),
    ("92","92-79.3 = +12.7","161.29"),
    ("68","68-79.3 = -11.3","127.69"),
    ("75","75-79.3 = -4.3","18.49"),
    ("88","88-79.3 = +8.7","75.69"),
    ("80","80-79.3 = +0.7","0.49"),
    ("SUM","","790.10"),
]
t2 = doc.add_table(rows=len(rows_data), cols=3)
t2.style = 'Table Grid'
for i,(a,b,c) in enumerate(rows_data):
    t2.rows[i].cells[0].text = a
    t2.rows[i].cells[1].text = b
    t2.rows[i].cells[2].text = c
    for cell in t2.rows[i].cells:
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    if i==0 or i==len(rows_data)-1:
        for cell in t2.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()
add_para(doc, "Step 2:")
add_formula(doc, "Sum[(xi - x_bar)^2] = 790.10")
add_para(doc, "Step 3:")
add_formula(doc, "s^2 = 790.10 / 9 = 87.79")
add_answer(doc, "Sample Variance = 87.79")
add_para(doc, "Note: NumPy returns 87.79 (exact: 87.7889) using precise floating-point arithmetic.", italic=True, color=(80,80,80))
add_para(doc, "Explanation: Bessel's correction (n-1) provides an unbiased estimate of population variance. Higher variance = more spread.", italic=True, color=(80,80,80))

add_separator(doc)

# STANDARD DEVIATION
add_para(doc, "6. SAMPLE STANDARD DEVIATION (n-1)", bold=True, size=12, color=(26,26,78))
add_para(doc, "Formula:")
add_formula(doc, "s = sqrt(s^2) = sqrt(87.79)")
add_formula(doc, "s = 9.37")
add_answer(doc, "Sample Standard Deviation = 9.37  (NumPy exact: 9.3696)")
add_para(doc, "Explanation: SD measures average deviation from mean. Students' marks deviate on average ~9.37 marks from 79.3.", italic=True, color=(80,80,80))

doc.add_page_break()

# ── Q 1.2 PROBABILITY ─────────────────────────────────────────────────────────
add_heading(doc, "Question 1.2 — Probability", level=2)

add_para(doc, "Q1 — Fair Coin Tossed Three Times: P(Exactly 2 Heads)", bold=True, size=12, color=(26,26,78))
add_para(doc, "Formula (Binomial):")
add_formula(doc, "P(X=k) = C(n,k) * p^k * (1-p)^(n-k)")
add_para(doc, "Given: n=3, k=2, p=0.5")
add_para(doc, "Step 1 — Sample Space (2^3 = 8 outcomes):")
add_formula(doc, "S = {HHH, HHT, HTH, HTT, THH, THT, TTH, TTT}")
add_para(doc, "Step 2 — Favourable outcomes (exactly 2 heads):")
add_formula(doc, "E = {HHT, HTH, THH}  =>  n(E) = 3")
add_para(doc, "Step 3 — Apply formula:")
add_formula(doc, "P(X=2) = C(3,2) * (0.5)^2 * (0.5)^1 = 3 * 0.25 * 0.5 = 3 * 0.125")
add_answer(doc, "P(Exactly 2 Heads) = 3/8 = 0.375 = 37.5%")
add_para(doc, "Explanation: 3 of 8 equally likely outcomes have exactly 2 heads. Probability = 3/8 = 37.5%.", italic=True, color=(80,80,80))

add_separator(doc)

add_para(doc, "Q2 — Balls in a Bag", bold=True, size=12, color=(26,26,78))
add_para(doc, "Given: Red=5, Blue=4, Green=3  =>  Total=12 balls")
add_para(doc, "Formula: P(E) = n(E) / n(S)")
doc.add_paragraph()
add_para(doc, "Part (a) — P(Red):")
add_formula(doc, "P(Red) = 5/12 = 0.4167")
add_answer(doc, "P(Red) = 5/12 ≈ 0.4167 = 41.67%")
doc.add_paragraph()
add_para(doc, "Part (b) — P(Not Green):")
add_formula(doc, "P(Green) = 3/12 = 1/4")
add_formula(doc, "P(Not Green) = 1 - 3/12 = 9/12 = 3/4")
add_answer(doc, "P(Not Green) = 3/4 = 0.75 = 75%")
add_para(doc, "Explanation: Complement rule: P(not A) = 1 - P(A). Non-green = 9 balls. P = 9/12 = 75%.", italic=True, color=(80,80,80))

add_separator(doc)

add_para(doc, "Q3 — Two Dice Rolled: P(Sum > 8)", bold=True, size=12, color=(26,26,78))
add_formula(doc, "Total sample space: n(S) = 6 x 6 = 36 outcomes")
add_para(doc, "Favourable outcomes (sum = 9, 10, 11, 12):")
fav_rows = [
    ("Sum","Favourable Outcomes","Count"),
    ("9","(3,6),(4,5),(5,4),(6,3)","4"),
    ("10","(4,6),(5,5),(6,4)","3"),
    ("11","(5,6),(6,5)","2"),
    ("12","(6,6)","1"),
    ("TOTAL","","10"),
]
t3 = doc.add_table(rows=len(fav_rows), cols=3)
t3.style = 'Table Grid'
for i,(a,b,c) in enumerate(fav_rows):
    t3.rows[i].cells[0].text = a
    t3.rows[i].cells[1].text = b
    t3.rows[i].cells[2].text = c
    for cell in t3.rows[i].cells:
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    if i==0 or i==len(fav_rows)-1:
        for cell in t3.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True
doc.add_paragraph()
add_formula(doc, "P(Sum > 8) = 10/36 = 5/18 ≈ 0.2778")
add_answer(doc, "P(Sum > 8) = 10/36 = 5/18 ≈ 0.2778 = 27.78%")
add_para(doc, "Explanation: 10 of 36 equally likely outcomes have sum > 8. Probability ≈ 27.78%.", italic=True, color=(80,80,80))

doc.add_page_break()

# ── Q 1.3 PERMUTATION & COMBINATION ──────────────────────────────────────────
add_heading(doc, "Question 1.3 — Permutation & Combination", level=2)

add_para(doc, "Problem 1 — Arrange letters of MACHINE (all vowels together)", bold=True, size=12, color=(26,26,78))
add_para(doc, "Word: M-A-C-H-I-N-E  |  Total letters = 7 (all distinct)")
add_para(doc, "Step 1 — Identify vowels: A, I, E (3 vowels)  |  Consonants: M, C, H, N (4)")
add_para(doc, "Step 2 — Treat vowels [AIE] as ONE block => 5 units: {[AIE], M, C, H, N}")
add_para(doc, "Step 3 — Arrange 5 units:")
add_formula(doc, "5! = 5 x 4 x 3 x 2 x 1 = 120 ways")
add_para(doc, "Step 4 — Arrange 3 vowels within the block:")
add_formula(doc, "3! = 3 x 2 x 1 = 6 ways")
add_para(doc, "Step 5 — Total (Multiplication Principle):")
add_formula(doc, "Total = 5! x 3! = 120 x 6 = 720")
add_answer(doc, "Total arrangements = 720")
add_para(doc, "Explanation: 5! = 120 ways to arrange the 5 units. 3! = 6 internal vowel arrangements. Total = 720.", italic=True, color=(80,80,80))

add_separator(doc)

add_para(doc, "Problem 2 — Committee of 4: Exactly 2 Men from 6, Exactly 2 Women from 5", bold=True, size=12, color=(26,26,78))
add_para(doc, "Formula: C(n,r) = n! / (r! * (n-r)!)")
add_para(doc, "Step 1 — Choose 2 men from 6:")
add_formula(doc, "C(6,2) = 6! / (2! x 4!) = (6 x 5) / (2 x 1) = 30/2 = 15")
add_para(doc, "Step 2 — Choose 2 women from 5:")
add_formula(doc, "C(5,2) = 5! / (2! x 3!) = (5 x 4) / (2 x 1) = 20/2 = 10")
add_para(doc, "Step 3 — Multiply (independent selections):")
add_formula(doc, "Total = C(6,2) x C(5,2) = 15 x 10 = 150")
add_answer(doc, "Total number of committees = 150")
add_para(doc, "Explanation: 15 ways to pick 2 men, 10 ways to pick 2 women. By multiplication rule: 15 x 10 = 150.", italic=True, color=(80,80,80))

doc.add_page_break()

# ── Q 1.4 BINOMIAL DISTRIBUTION ───────────────────────────────────────────────
add_heading(doc, "Question 1.4 — Binomial Distribution", level=2)
add_para(doc, "Given: Biased coin P(Head)=0.6, P(Tail)=0.4, n=5 tosses", bold=True)
add_para(doc, "Formula: P(X=k) = C(n,k) x p^k x (1-p)^(n-k)", bold=True)

add_separator(doc)
add_para(doc, "Part (a) — P(Exactly 3 Heads) = P(X=3)", bold=True, size=12, color=(26,26,78))
add_para(doc, "n=5, k=3, p=0.6, (1-p)=0.4")
add_para(doc, "Step 1:")
add_formula(doc, "C(5,3) = 5! / (3! x 2!) = (5 x 4) / (2 x 1) = 10")
add_para(doc, "Step 2:")
add_formula(doc, "(0.6)^3 = 0.6 x 0.6 x 0.6 = 0.216")
add_para(doc, "Step 3:")
add_formula(doc, "(0.4)^2 = 0.4 x 0.4 = 0.16")
add_para(doc, "Step 4:")
add_formula(doc, "P(X=3) = 10 x 0.216 x 0.16 = 10 x 0.03456")
add_answer(doc, "P(X=3) = 0.3456 = 34.56%")
add_para(doc, "Explanation: C(5,3)=10 counts all orderings of 3H in 5 tosses. Combined probability = 34.56%.", italic=True, color=(80,80,80))

add_separator(doc)
add_para(doc, "Part (b) — P(At Least 4 Heads) = P(X>=4) = P(X=4) + P(X=5)", bold=True, size=12, color=(26,26,78))

add_para(doc, "Calculate P(X=4):")
add_formula(doc, "C(5,4) = 5,  (0.6)^4 = 0.1296,  (0.4)^1 = 0.4")
add_formula(doc, "P(X=4) = 5 x 0.1296 x 0.4 = 5 x 0.05184 = 0.2592")

add_para(doc, "Calculate P(X=5):")
add_formula(doc, "C(5,5) = 1,  (0.6)^5 = 0.07776,  (0.4)^0 = 1")
add_formula(doc, "P(X=5) = 1 x 0.07776 x 1 = 0.07776")

add_para(doc, "Final Calculation:")
add_formula(doc, "P(X>=4) = P(X=4) + P(X=5) = 0.2592 + 0.07776 = 0.33696")
add_answer(doc, "P(X>=4) = 0.3370 = 33.70%")
add_para(doc, "Explanation: Sum the probabilities for 4 and 5 heads. Since p=0.6>0.5, higher heads are common: 33.70%.", italic=True, color=(80,80,80))

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# TASK 2 — PYTHON
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "TASK 2 — Python, NumPy & Pandas", level=1, color=(26,26,78))
add_para(doc, "All code is contained in: Student_Performance_Assignment1.ipynb", bold=True)
add_para(doc, "Below shows the code and expected outputs for each section.", italic=True)
doc.add_paragraph()

add_para(doc, "Task 2.1 — NumPy Statistical Verification", bold=True, size=12, color=(26,26,78))
code1 = """import numpy as np
marks = np.array([72, 85, 90, 65, 78, 92, 68, 75, 88, 80])
print(f"Mean     : {np.mean(marks):.4f}")
print(f"Median   : {np.median(marks):.4f}")
print(f"Range    : {int(np.max(marks)-np.min(marks))}")
print(f"Variance : {np.var(marks, ddof=1):.4f}")
print(f"Std Dev  : {np.std(marks, ddof=1):.4f}")"""
p=doc.add_paragraph()
run=p.add_run(code1)
run.font.name="Courier New"
run.font.size=Pt(9)

add_para(doc, "Expected Output:", bold=True, color=(0,100,0))
output1="""Mean     : 79.3000
Median   : 79.0000
Range    : 27
Variance : 87.7889  (sample, ddof=1)
Std Dev  : 9.3696   (sample, ddof=1)"""
p=doc.add_paragraph()
run=p.add_run(output1)
run.font.name="Courier New"
run.font.size=Pt(9)
run.font.color.rgb=RGBColor(0,100,0)

add_separator(doc)

add_para(doc, "Task 2.2 — Average Salary by Department (groupby)", bold=True, size=12, color=(26,26,78))
add_para(doc, "groupby() Result:")
dept_table=[
    ("Department","Average Salary (Rs.)"),
    ("Finance","78,500.00"),
    ("IT","61,250.00"),
    ("Sales","50,000.00"),
    ("HR","45,000.00"),
]
t4=doc.add_table(rows=len(dept_table),cols=2)
t4.style='Table Grid'
for i,(a,b) in enumerate(dept_table):
    t4.rows[i].cells[0].text=a
    t4.rows[i].cells[1].text=b
    for cell in t4.rows[i].cells:
        cell.paragraphs[0].runs[0].font.size=Pt(10)
    if i==0:
        for cell in t4.rows[i].cells:
            cell.paragraphs[0].runs[0].bold=True

doc.add_paragraph()
add_para(doc, "Employees with Experience > 5: Diya, Gaurav, Ishan, Jiya (4 employees)", bold=False)
add_para(doc, "Top 3 Salaries: Gaurav (Rs.85,000), Jiya (Rs.78,000), Diya (Rs.72,000)")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# TASK 3 — HYPOTHESIS TESTING
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "TASK 3 — Hypothesis Testing", level=1, color=(26,26,78))

add_para(doc, "Working Hours: 7.5, 8.2, 7.8, 9.0, 8.5, 7.2, 8.8, 9.1, 7.4, 8.0", bold=True)
add_para(doc, "Company Claim: Average = 8 hours  |  Significance Level: alpha = 0.05", bold=True)

add_heading(doc, "Task 3.1 — One-Sample t-Test", level=2)

add_para(doc, "H0: mu = 8  (Null: Mean working hours = 8)", bold=True)
add_para(doc, "H1: mu != 8  (Alternative: Mean is not 8, Two-tailed)", bold=True)
doc.add_paragraph()

steps_ht = [
    ("Sample Mean",    "x_bar = (7.5+8.2+7.8+9.0+8.5+7.2+8.8+9.1+7.4+8.0)/10 = 81.5/10", "x_bar = 8.15"),
    ("Sample SD",      "s = sqrt(4.2050/9) = sqrt(0.4672)",                                  "s = 0.6835"),
    ("Degrees of Freedom","df = n - 1 = 10 - 1",                                             "df = 9"),
    ("t-Statistic",    "t = (8.15 - 8.0) / (0.6835/sqrt(10)) = 0.15 / 0.2161",             "t = 0.6941"),
    ("t-Critical",     "t(0.025, 9) from t-table [two-tailed, alpha=0.05]",                  "t_crit = +/-2.262"),
    ("p-Value",        "Computed using SciPy: stats.ttest_1samp()",                          "p = 0.5052"),
    ("Decision",       "|t_calc|=0.6941 < t_crit=2.262",                                    "FAIL TO REJECT H0"),
]
t5=doc.add_table(rows=len(steps_ht)+1, cols=3)
t5.style='Table Grid'
header_row=["Step","Calculation","Result"]
for j,h in enumerate(header_row):
    t5.rows[0].cells[j].text=h
    t5.rows[0].cells[j].paragraphs[0].runs[0].bold=True
    t5.rows[0].cells[j].paragraphs[0].runs[0].font.size=Pt(10)
for i,(a,b,c) in enumerate(steps_ht):
    t5.rows[i+1].cells[0].text=a
    t5.rows[i+1].cells[1].text=b
    t5.rows[i+1].cells[2].text=c
    for cell in t5.rows[i+1].cells:
        cell.paragraphs[0].runs[0].font.size=Pt(10)
    t5.rows[i+1].cells[0].paragraphs[0].runs[0].bold=True

doc.add_paragraph()
add_para(doc, "CONCLUSION: At alpha=0.05, there is insufficient evidence to reject H0. The company's claim that average working hours = 8 is statistically supported. The t-statistic (0.6941) lies well within the acceptance region (-2.262 to +2.262).", bold=False, italic=True, color=(0,80,0))

add_separator(doc)

add_heading(doc, "Task 3.2 — 95% Confidence Interval", level=2)
add_para(doc, "Formula: CI = x_bar +/- t(alpha/2, df) * (s/sqrt(n))")
add_para(doc, "Given: x_bar=8.15, s=0.6835, n=10, t_crit=2.262")
doc.add_paragraph()
add_formula(doc, "Margin of Error = 2.262 x (0.6835/sqrt(10)) = 2.262 x 0.2161 = 0.4888")
add_formula(doc, "Lower Limit (LL) = 8.15 - 0.4888 = 7.6612")
add_formula(doc, "Upper Limit (UL) = 8.15 + 0.4888 = 8.6388")
add_answer(doc, "95% Confidence Interval = (7.66, 8.64) hours")
add_para(doc, "Interpretation: We are 95% confident the true mean lies between 7.66 and 8.64 hours. Since mu_0=8 is inside this interval, it is consistent with failing to reject H0.", italic=True, color=(80,80,80))

add_separator(doc)

add_heading(doc, "Task 3.3 — Matplotlib Visualizations", level=2)
add_para(doc, "Three professional charts were generated and saved:", bold=True)
charts=[
    ("bar_chart.png",    "Bar Chart",    "Department vs Average Salary — shows Finance dept has highest avg salary (Rs.78,500)"),
    ("scatter_plot.png", "Scatter Plot", "Experience vs Salary — positive correlation visible with trend line"),
    ("histogram.png",    "Histogram",    "Age Distribution — employees distributed from 26-40, mean age = 31.8 yrs"),
]
t6=doc.add_table(rows=len(charts)+1, cols=3)
t6.style='Table Grid'
for j,h in enumerate(["File","Chart Type","Description"]):
    t6.rows[0].cells[j].text=h
    t6.rows[0].cells[j].paragraphs[0].runs[0].bold=True
    t6.rows[0].cells[j].paragraphs[0].runs[0].font.size=Pt(10)
for i,(a,b,c) in enumerate(charts):
    t6.rows[i+1].cells[0].text=a
    t6.rows[i+1].cells[1].text=b
    t6.rows[i+1].cells[2].text=c
    for cell in t6.rows[i+1].cells:
        cell.paragraphs[0].runs[0].font.size=Pt(9)

doc.add_paragraph()

# Embed charts if they exist
for fname in ["bar_chart.png", "scatter_plot.png", "histogram.png"]:
    if os.path.exists(fname):
        add_para(doc, fname, bold=True)
        doc.add_picture(fname, width=Inches(5.5))
        doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SUMMARY TABLE
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "Assignment 1 — Complete Summary", level=1, color=(26,26,78))

summary=[
    ("Task","Topic","Key Result","Status"),
    ("1.1","Descriptive Statistics","Mean=79.3, Median=79, SD=9.37","Done"),
    ("1.2 Q1","Coin Toss","P(2 Heads)=3/8=0.375","Done"),
    ("1.2 Q2","Balls in Bag","P(Red)=5/12, P(Not Green)=3/4","Done"),
    ("1.2 Q3","Two Dice","P(Sum>8)=5/18=0.2778","Done"),
    ("1.3 P1","MACHINE Vowels","720 arrangements","Done"),
    ("1.3 P2","Committee","C(6,2)xC(5,2)=150","Done"),
    ("1.4 (a)","Binomial P(X=3)","0.3456 (34.56%)","Done"),
    ("1.4 (b)","Binomial P(X>=4)","0.3370 (33.70%)","Done"),
    ("2.1","NumPy Stats","All verified","Done"),
    ("2.2","Pandas Ops","8 operations completed","Done"),
    ("3.1","Hypothesis Test","Fail to Reject H0","Done"),
    ("3.2","Confidence Interval","(7.66, 8.64) hours","Done"),
    ("3.3","Matplotlib Charts","3 charts saved as PNG","Done"),
]
t7=doc.add_table(rows=len(summary),cols=4)
t7.style='Table Grid'
for i,row in enumerate(summary):
    for j,val in enumerate(row):
        t7.rows[i].cells[j].text=val
        run=t7.rows[i].cells[j].paragraphs[0].runs[0]
        run.font.size=Pt(10)
        if i==0:
            run.bold=True
        if j==3 and i>0:
            run.font.color.rgb=RGBColor(0,128,0)
            run.bold=True

doc.add_paragraph()
p=doc.add_paragraph()
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=p.add_run("Student: Bablu Kumar  |  Roll No.: 23052440012  |  Date: 09/05/26")
run.bold=True
run.font.size=Pt(11)
run.font.color.rgb=RGBColor(26,26,78)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save("Assignment_1_Solution.docx")
print("Assignment_1_Solution.docx saved successfully.")
